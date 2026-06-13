"""Shared utilities for generating Word documents (phiếu) from warehouse transactions."""

from __future__ import annotations


# ── Vietnamese number → words ─────────────────────────────────────────────────

def so_thanh_chu(amount: int) -> str:
    if amount == 0:
        return "Không đồng"
    _U = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]

    def _b(n):
        t, c, d = n // 100, (n % 100) // 10, n % 10
        p = []
        if t:
            p.append(_U[t] + " trăm")
        if c == 1:
            p.append("mười")
            if d == 5:   p.append("lăm")
            elif d:      p.append(_U[d])
        elif c > 1:
            p.append(_U[c] + " mươi")
            if d == 1:   p.append("mốt")
            elif d == 5: p.append("lăm")
            elif d:      p.append(_U[d])
        elif t and d:
            p.append("lẻ " + _U[d])
        elif d:
            p.append(_U[d])
        return " ".join(p)

    nghin_ty = amount // 1_000_000_000_000
    ty       = (amount % 1_000_000_000_000) // 1_000_000_000
    trieu    = (amount % 1_000_000_000) // 1_000_000
    nghin    = (amount % 1_000_000) // 1_000
    tram     = amount % 1_000
    parts  = []
    if nghin_ty: parts.append(_b(nghin_ty % 1000) + " nghìn tỷ")
    if ty:       parts.append(_b(ty)    + " tỷ")
    if trieu:    parts.append(_b(trieu) + " triệu")
    if nghin:    parts.append(_b(nghin) + " nghìn")
    if tram:     parts.append(_b(tram))
    s = " ".join(parts)
    return s[0].upper() + s[1:] + " đồng"


def _parse_date(s: str):
    from datetime import datetime
    for fmt in ("%Y-%m-%d", "%d/%m/%Y"):
        try:
            dt = datetime.strptime(s, fmt)
            return f"{dt.day:02d}", f"{dt.month:02d}", str(dt.year)
        except ValueError:
            pass
    return "___", "___", "______"


# ── Document skeleton ─────────────────────────────────────────────────────────

def _new_doc():
    from docx import Document
    from docx.shared import Cm
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = sec.page_height = None  # reset first
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.0)
    return doc


def _helpers():
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FN = "Times New Roman"
    L, C, R = WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT

    def _r(para, text, sz=12, bold=False, italic=False, ul=False):
        rn = para.add_run(text)
        rn.font.name = FN
        rn.font.size = Pt(sz)
        rn.font.bold, rn.font.italic, rn.font.underline = bold, italic, ul
        return rn

    def _fp(p, after=0, before=0, align=None):
        p.paragraph_format.space_after  = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        if align is not None:
            p.alignment = align

    def _ct(cell, text, sz=11, bold=False, italic=False, align=C, ul=False):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.clear()
        _fp(p, after=1, before=1, align=align)
        _r(p, text, sz, bold, italic, ul)

    def _set_bdr(cell, solid=True, sz=6):
        tcPr = cell._tc.get_or_add_tcPr()
        bdr  = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single" if solid else "none")
            if solid:
                el.set(qn("w:sz"), str(sz))
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
            bdr.append(el)
        tcPr.append(bdr)

    def _borders(table, solid=True):
        seen = set()
        for row in table.rows:
            for cell in row.cells:
                cid = id(cell._tc)
                if cid not in seen:
                    seen.add(cid)
                    _set_bdr(cell, solid)

    return _r, _fp, _ct, _set_bdr, _borders, L, C, R


# ── Shared layout blocks ───────────────────────────────────────────────────────

def _title_section(doc, h, title, ref_number,
                   unit1="QUÂN KHU 5", unit2="CỤC HẬU CẦN – KỸ THUẬT"):
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    t = doc.add_table(rows=1, cols=3)
    _borders(t, solid=False)

    lc = t.cell(0, 0)
    p1 = lc.paragraphs[0]
    _fp(p1, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    _r(p1, unit1, 12)
    p2 = lc.add_paragraph()
    _fp(p2, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    _r(p2, unit2, 12, bold=True, ul=True)

    _ct(t.cell(0, 1), title, sz=16, bold=True)
    _ct(t.cell(0, 2), f"Số phiếu: {ref_number}", sz=12, align=R)

    t.columns[0].width = Cm(5.5)
    t.columns[1].width = Cm(8.0)
    t.columns[2].width = Cm(4.5)


def _info_section(doc, h, left_lines, notes):
    """left_lines: list of (label, value) tuples."""
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h
    from docx.shared import Cm

    t = doc.add_table(rows=1, cols=2)
    _borders(t, solid=False)

    lc = t.cell(0, 0)
    for i, (lbl, val) in enumerate(left_lines):
        p = lc.paragraphs[0] if i == 0 else lc.add_paragraph()
        _fp(p, after=2)
        _r(p, f"{lbl}: ", 12)
        _r(p, val or "", 12)

    rc = t.cell(0, 1)
    rp = rc.paragraphs[0]
    _fp(rp, after=2)
    _r(rp, "Nội dung: ", 12)
    _r(rp, notes or "", 12)

    t.columns[0].width = Cm(8.5)
    t.columns[1].width = Cm(9.5)


def _table_with_split_qty(doc, h, lines, show_price=True):
    """Table: TT | Tên hàng | ĐVT | Kế hoạch | Thực hiện | Đơn giá | Thành tiền | Ghi chú"""
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h
    from docx.shared import Cm

    n = len(lines)
    tbl = doc.add_table(rows=2 + n + 1, cols=8)

    tbl.cell(0, 3).merge(tbl.cell(0, 4))           # "Số lượng" across 2 cols
    for ci in [0, 1, 2, 5, 6, 7]:
        tbl.cell(0, ci).merge(tbl.cell(1, ci))     # span 2 rows
    tbl.cell(n + 2, 0).merge(tbl.cell(n + 2, 5))  # "Tổng cộng" span

    _borders(tbl, solid=True)

    _ct(tbl.cell(0, 0), "TT",          bold=True)
    _ct(tbl.cell(0, 1), "Tên hàng",    bold=True)
    _ct(tbl.cell(0, 2), "ĐVT",         bold=True)
    _ct(tbl.cell(0, 3), "Số lượng",    bold=True)
    _ct(tbl.cell(1, 3), "Kế\nhoạch",   bold=True)
    _ct(tbl.cell(1, 4), "Thực\nhiện",  bold=True)
    _ct(tbl.cell(0, 5), "Đơn giá",     bold=True)
    _ct(tbl.cell(0, 6), "Thành tiền",  bold=True)
    _ct(tbl.cell(0, 7), "Ghi chú",     bold=True)

    total = 0
    for i, line in enumerate(lines):
        ri  = 2 + i
        pr  = line.unit_price if show_price else 0
        lt  = int(line.quantity * pr)
        total += lt
        _ct(tbl.cell(ri, 0), f"{i+1:02d}", align=C)
        _ct(tbl.cell(ri, 1), line.item_name, align=L)
        _ct(tbl.cell(ri, 2), line.unit_of_measure, align=C)
        _ct(tbl.cell(ri, 3), str(line.quantity), align=C)
        _ct(tbl.cell(ri, 4), "", align=C)
        _ct(tbl.cell(ri, 5), f"{pr:,.0f}" if pr else "", align=R)
        _ct(tbl.cell(ri, 6), f"{lt:,.0f}" if lt else "", align=R)
        _ct(tbl.cell(ri, 7), line.notes or "", align=L)

    _ct(tbl.cell(n + 2, 0), "Tổng cộng", bold=True)
    _ct(tbl.cell(n + 2, 6), f"{total:,.0f}" if total else "", bold=True, align=R)
    _ct(tbl.cell(n + 2, 7), "")

    for i, w in enumerate([1.0, 4.5, 1.2, 1.4, 1.4, 1.8, 2.0, 1.7]):
        for row in tbl.rows:
            row.cells[i].width = Cm(w)

    return total


def _table_simple(doc, h, lines, show_quality=True):
    """Table: TT | Tên hàng | ĐVT | Số lượng | [Mức HH] | Ghi chú"""
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h
    from docx.shared import Cm

    n      = len(lines)
    n_cols = 6 if show_quality else 5
    tbl    = doc.add_table(rows=1 + n + 1, cols=n_cols)
    tbl.cell(n + 1, 0).merge(tbl.cell(n + 1, n_cols - 2))

    _borders(tbl, solid=True)

    _ct(tbl.cell(0, 0), "TT",       bold=True)
    _ct(tbl.cell(0, 1), "Tên hàng", bold=True)
    _ct(tbl.cell(0, 2), "ĐVT",      bold=True)
    _ct(tbl.cell(0, 3), "Số lượng", bold=True)
    if show_quality:
        _ct(tbl.cell(0, 4), "Mức HH", bold=True)
    _ct(tbl.cell(0, n_cols - 1), "Ghi chú", bold=True)

    for i, line in enumerate(lines):
        ri = 1 + i
        _ct(tbl.cell(ri, 0), f"{i+1:02d}", align=C)
        _ct(tbl.cell(ri, 1), line.item_name, align=L)
        _ct(tbl.cell(ri, 2), line.unit_of_measure, align=C)
        _ct(tbl.cell(ri, 3), str(line.quantity), align=C)
        if show_quality:
            _ct(tbl.cell(ri, 4), getattr(line, "quality_level", ""), align=C)
        _ct(tbl.cell(ri, n_cols - 1), line.notes or "", align=L)

    _ct(tbl.cell(n + 1, 0), "Tổng cộng", bold=True)
    _ct(tbl.cell(n + 1, n_cols - 1), "")

    widths = [1.0, 5.5, 1.2, 1.5, 1.3, 2.5] if show_quality else [1.0, 6.5, 1.2, 1.5, 2.8]
    for i, w in enumerate(widths[:n_cols]):
        for row in tbl.rows:
            row.cells[i].width = Cm(w)


def _footer_section(doc, h, n_items, total, date_str, show_price, sig_labels=None):
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if sig_labels is None:
        sig_labels = ["NGƯỜI LẬP PHIẾU", "NGƯỜI GIAO", "NGƯỜI NHẬN",
                      "PHÒNG QUÂN NHU", "THỦ TRƯỞNG"]

    p = doc.add_paragraph()
    _fp(p, after=2)
    _r(p, f"Tổng:  {n_items:02d} Khoản", 12)

    if show_price and total:
        pw = doc.add_paragraph()
        _fp(pw, after=2)
        _r(pw, "Thành tiền: ", 12)
        _r(pw, so_thanh_chu(int(total)) + ".", 12)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    d, m, y = _parse_date(date_str)
    pd = doc.add_paragraph()
    _fp(pd, after=6, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _r(pd, f"Ngày {d}  tháng {m}  năm {y}", 12, italic=True)

    ts = doc.add_table(rows=2, cols=len(sig_labels))
    _borders(ts, solid=False)
    for ci, lbl in enumerate(sig_labels):
        _ct(ts.cell(0, ci), lbl, bold=True)
        _ct(ts.cell(1, ci), " " * 16)


def _save_path(parent, ref_number, prefix):
    from PyQt6.QtWidgets import QFileDialog
    name = f"{prefix}_{ref_number or 'export'}.docx"
    path, _ = QFileDialog.getSaveFileName(
        parent, "Lưu Phiếu Word", name, "Word Files (*.docx)"
    )
    return path


def _ensure_docx(parent) -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.warning(parent, "Thiếu thư viện",
                            "Vui lòng cài đặt:\n\npip install python-docx")
        return False


# ── Public export functions ────────────────────────────────────────────────────

def export_nhap_kho(parent, receipt, lines) -> None:
    if not _ensure_docx(parent):
        return
    from PyQt6.QtWidgets import QMessageBox
    from docx.shared import Pt

    path = _save_path(parent, receipt.reference_number, "PhieuNhap")
    if not path:
        return

    show_price   = (receipt.subtype == "new")
    show_quality = (receipt.subtype == "from_unit")

    doc = _new_doc()
    h   = _helpers()
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h

    _title_section(doc, h, "PHIẾU NHẬP KHO", receipt.reference_number)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    source   = receipt.from_warehouse_name or receipt.supplier or ""
    src_lbl  = {
        "from_unit":    "Đơn vị giao",
        "unit_return":  "Đơn vị trả",
        "event_return": "Nguồn trả",
    }.get(receipt.subtype, "Đơn vị giao")

    left = [("Giao tại kho", receipt.to_warehouse_name), (src_lbl, source)]
    if receipt.subtype == "new" and receipt.transporter:
        left.append(("Hàng do", receipt.transporter))

    _info_section(doc, h, left, receipt.notes)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    if show_price:
        total = _table_with_split_qty(doc, h, lines, show_price=True)
    else:
        _table_simple(doc, h, lines, show_quality=show_quality)
        total = 0

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _footer_section(doc, h, len(lines), total, receipt.transaction_date, show_price)

    doc.save(path)
    QMessageBox.information(parent, "Xuất Word", f"Đã lưu:\n{path}")


def export_nhap_kho_excel(parent, receipt, lines) -> None:
    """Xuất phiếu nhập kho ra Excel theo mẫu chuẩn (8 cột, header 2 dòng)."""
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    except ImportError:
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.warning(parent, "Lỗi", "Cần cài openpyxl:\npip install openpyxl")
        return

    from PyQt6.QtWidgets import QFileDialog, QMessageBox

    path, _ = QFileDialog.getSaveFileName(
        parent, "Lưu Phiếu Nhập Kho",
        f"PhieuNhapKho_{receipt.reference_number or 'export'}.xlsx",
        "Excel Files (*.xlsx)",
    )
    if not path:
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Phiếu Nhập Kho"
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = 9  # A4

    TN = "Times New Roman"
    thin = Side(style="thin", color="000000")
    bdr = Border(top=thin, left=thin, right=thin, bottom=thin)

    def aln(h="left", v="center", wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    def sh(row, height):
        ws.row_dimensions[row].height = height

    for col, w in [("A", 5), ("B", 28), ("C", 7), ("D", 10), ("E", 10),
                   ("F", 12), ("G", 14), ("H", 18)]:
        ws.column_dimensions[col].width = w

    # ── Row 1: QUÂN KHU 5 | Số phiếu ────────────────────────────────────────
    sh(1, 18)
    ws.merge_cells("A1:B1")
    ws["A1"].value = "QUÂN KHU 5"
    ws["A1"].font = Font(name=TN, size=11)
    ws["A1"].alignment = aln("center")
    ws.merge_cells("G1:H1")
    ws["G1"].value = f"Số phiếu: {receipt.reference_number or ''}"
    ws["G1"].font = Font(name=TN, size=11)
    ws["G1"].alignment = aln("right")

    # ── Row 2: CỤC HẬU CẦN – KỸ THUẬT | PHIẾU NHẬP KHO ─────────────────────
    sh(2, 26)
    ws.merge_cells("A2:B2")
    ws["A2"].value = "CỤC HẬU CẦN – KỸ THUẬT"
    ws["A2"].font = Font(name=TN, size=11, bold=True, underline="single")
    ws["A2"].alignment = aln("center")
    ws.merge_cells("C2:F2")
    ws["C2"].value = "PHIẾU NHẬP KHO"
    ws["C2"].font = Font(name=TN, size=16, bold=True)
    ws["C2"].alignment = aln("center")

    # ── Row 3: blank ──────────────────────────────────────────────────────────
    sh(3, 8)

    # ── Rows 4-6: info block ─────────────────────────────────────────────────
    source  = receipt.from_warehouse_name or receipt.supplier or ""
    src_lbl = {"from_unit": "Đơn vị giao", "unit_return": "Đơn vị trả",
               "event_return": "Nguồn trả"}.get(receipt.subtype, "Đơn vị giao")
    left_info = [
        f"Giao tại kho: {receipt.to_warehouse_name or ''}",
        f"{src_lbl}: {source}",
        f"Hàng do: {receipt.transporter or ''}",
    ]
    for ri, text in enumerate(left_info, start=4):
        sh(ri, 18)
        ws.merge_cells(f"A{ri}:C{ri}")
        c = ws.cell(row=ri, column=1, value=text)
        c.font = Font(name=TN, size=11)
        c.alignment = aln("left")
    ws.merge_cells("D4:H6")
    ws["D4"].value = f"Nội dung: {receipt.notes or ''}"
    ws["D4"].font = Font(name=TN, size=11)
    ws["D4"].alignment = aln("left", wrap=True)

    # ── Row 7: blank ──────────────────────────────────────────────────────────
    sh(7, 8)

    # ── Rows 8-9: table header (double row) ──────────────────────────────────
    hf = Font(name=TN, size=11, bold=True)
    ha = Alignment(horizontal="center", vertical="center", wrap_text=True)
    hfill = PatternFill("solid", fgColor="F2F2F2")
    sh(8, 22); sh(9, 20)
    for ci, val in enumerate(["TT","Tên hàng","ĐVT","Số lượng","","Đơn giá","Thành tiền","Ghi chú"], 1):
        c = ws.cell(row=8, column=ci, value=val)
        c.font = hf; c.alignment = ha; c.border = bdr; c.fill = hfill
    for ci, val in enumerate(["","","","Kế hoạch","Thực hiện","","",""], 1):
        c = ws.cell(row=9, column=ci, value=val)
        c.font = hf; c.alignment = ha; c.border = bdr; c.fill = hfill
    ws.merge_cells("A8:A9"); ws.merge_cells("B8:B9"); ws.merge_cells("C8:C9")
    ws.merge_cells("D8:E8")
    ws.merge_cells("F8:F9"); ws.merge_cells("G8:G9"); ws.merge_cells("H8:H9")

    # ── Data rows ─────────────────────────────────────────────────────────────
    grand_total = 0.0
    DR = 10
    for i, line in enumerate(lines):
        r = DR + i
        sh(r, 20)
        price = getattr(line, "unit_price", 0.0) or 0.0
        lt = line.quantity * price
        grand_total += lt
        vals = [i + 1, line.item_name, line.unit_of_measure, line.quantity, "",
                price or "", lt or "", getattr(line, "notes", "") or ""]
        alns = [aln("center"), aln("left"), aln("center"), aln("center"), aln("center"),
                aln("right"), aln("right"), aln("left")]
        for ci, (val, al) in enumerate(zip(vals, alns), 1):
            c = ws.cell(row=r, column=ci, value=val)
            c.font = Font(name=TN, size=11); c.alignment = al; c.border = bdr
            if ci in (6, 7) and isinstance(val, (int, float)) and val:
                c.number_format = "#,##0"

    # ── Tổng cộng row ─────────────────────────────────────────────────────────
    tr = DR + len(lines)
    sh(tr, 22)
    ws.merge_cells(f"A{tr}:F{tr}")
    c = ws.cell(row=tr, column=1, value="Tổng cộng")
    c.font = Font(name=TN, size=11, bold=True); c.alignment = aln("center"); c.border = bdr
    c = ws.cell(row=tr, column=7, value=grand_total if grand_total else "")
    c.font = Font(name=TN, size=11, bold=True); c.alignment = aln("right"); c.border = bdr
    if grand_total:
        c.number_format = "#,##0"
    ws.cell(row=tr, column=8).border = bdr

    # ── Footer ────────────────────────────────────────────────────────────────
    f1 = tr + 1; sh(f1, 18)
    ws.merge_cells(f"A{f1}:H{f1}")
    c = ws.cell(row=f1, column=1, value=f"Tổng: {len(lines):02d} Khoản")
    c.font = Font(name=TN, size=11); c.alignment = aln("left")

    f2 = tr + 2; sh(f2, 18)
    ws.merge_cells(f"A{f2}:H{f2}")
    c = ws.cell(row=f2, column=1,
                value=f"Thành tiền: {so_thanh_chu(int(grand_total))}.")
    c.font = Font(name=TN, size=11); c.alignment = aln("left")

    dy, mo, yr = _parse_date(receipt.transaction_date or "")
    dr = tr + 3; sh(dr, 18)
    ws.merge_cells(f"E{dr}:H{dr}")
    c = ws.cell(row=dr, column=5,
                value=f"Ngày {dy}  tháng {mo}  năm {yr}")
    c.font = Font(name=TN, size=11, italic=True); c.alignment = aln("center")

    sh(dr + 1, 14)

    sr = dr + 2; sh(sr, 18)
    sigs = [("A", "A", "NGƯỜI LẬP PHIẾU"), ("B", "C", "NGƯỜI GIAO"),
            ("D", "E", "NGƯỜI NHẬN"), ("F", "G", "PHÒNG QUÂN NHU"), ("H", "H", "THỦ TRƯỞNG")]
    for cs, ce, label in sigs:
        if cs != ce:
            ws.merge_cells(f"{cs}{sr}:{ce}{sr}")
        ci = ord(cs) - ord("A") + 1
        c = ws.cell(row=sr, column=ci, value=label)
        c.font = Font(name=TN, size=11, bold=True); c.alignment = aln("center")

    try:
        wb.save(path)
        QMessageBox.information(parent, "Xuất Excel",
                                f"Đã lưu phiếu nhập kho:\n{path}")
    except Exception as exc:
        QMessageBox.warning(parent, "Lỗi", f"Không thể lưu file:\n{exc}")


def export_xuat_kho(parent, issue, lines) -> None:
    if not _ensure_docx(parent):
        return
    from PyQt6.QtWidgets import QMessageBox
    from docx.shared import Pt

    path = _save_path(parent, issue.reference_number, "PhieuXuat")
    if not path:
        return

    show_price = (issue.subtype == "to_unit")

    doc = _new_doc()
    h   = _helpers()
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h

    _title_section(doc, h, "PHIẾU XUẤT KHO", issue.reference_number)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    dest     = issue.to_warehouse_name or issue.recipient or ""
    dest_lbl = "Đơn vị nhận" if issue.subtype == "to_unit" else "Đơn vị mượn"

    left = [("Xuất tại kho", issue.from_warehouse_name), (dest_lbl, dest)]
    if issue.subtype == "to_unit" and issue.transporter:
        left.append(("Hàng do", issue.transporter))

    _info_section(doc, h, left, issue.notes)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    if show_price:
        total = _table_with_split_qty(doc, h, lines, show_price=True)
    else:
        _table_simple(doc, h, lines, show_quality=True)
        total = 0

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    sig = ["NGƯỜI LẬP PHIẾU", "NGƯỜI GIAO", "NGƯỜI NHẬN", "PHÒNG QUÂN NHU", "THỦ TRƯỞNG"]
    _footer_section(doc, h, len(lines), total, issue.transaction_date, show_price, sig)

    doc.save(path)
    QMessageBox.information(parent, "Xuất Word", f"Đã lưu:\n{path}")


def export_xuat_kho_excel(parent, issue, lines) -> None:
    """Xuất phiếu xuất kho (đi đơn vị) ra Excel theo mẫu chuẩn."""
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    except ImportError:
        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.warning(parent, "Lỗi", "Cần cài openpyxl:\npip install openpyxl")
        return

    from PyQt6.QtWidgets import QFileDialog, QMessageBox
    from datetime import datetime, timedelta

    path, _ = QFileDialog.getSaveFileName(
        parent, "Lưu Phiếu Xuất Kho",
        f"PhieuXuatKho_{issue.reference_number or 'export'}.xlsx",
        "Excel Files (*.xlsx)",
    )
    if not path:
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Phiếu Xuất Kho"
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = 9  # A4

    TN = "Times New Roman"
    thin = Side(style="thin", color="000000")
    bdr  = Border(top=thin, left=thin, right=thin, bottom=thin)

    def aln(h="left", v="center", wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    def sh(row, height):
        ws.row_dimensions[row].height = height

    for col, w in [("A", 5), ("B", 28), ("C", 7), ("D", 10), ("E", 10),
                   ("F", 12), ("G", 14), ("H", 18)]:
        ws.column_dimensions[col].width = w

    # ── Row 1: QUÂN KHU 5 | Số phiếu ────────────────────────────────────────
    sh(1, 18)
    ws.merge_cells("A1:B1")
    ws["A1"].value = "QUÂN KHU 5"
    ws["A1"].font = Font(name=TN, size=11)
    ws["A1"].alignment = aln("center")
    ws.merge_cells("G1:H1")
    ws["G1"].value = f"Số phiếu: {issue.reference_number or ''}"
    ws["G1"].font = Font(name=TN, size=11)
    ws["G1"].alignment = aln("right")

    # ── Row 2: CỤC HẬU CẦN – KỸ THUẬT | PHIẾU XUẤT KHO ─────────────────────
    sh(2, 26)
    ws.merge_cells("A2:B2")
    ws["A2"].value = "CỤC HẬU CẦN – KỸ THUẬT"
    ws["A2"].font = Font(name=TN, size=11, bold=True, underline="single")
    ws["A2"].alignment = aln("center")
    ws.merge_cells("C2:F2")
    ws["C2"].value = "PHIẾU XUẤT KHO"
    ws["C2"].font = Font(name=TN, size=16, bold=True)
    ws["C2"].alignment = aln("center")

    # ── Row 3: blank ──────────────────────────────────────────────────────────
    sh(3, 8)

    # ── Rows 4-7: info block ─────────────────────────────────────────────────
    # "Có giá trị đến ngày" = transaction_date + 60 ngày
    try:
        tx_dt = datetime.strptime(issue.transaction_date, "%Y-%m-%d")
        valid_until = (tx_dt + timedelta(days=60)).strftime("%d/%m/%Y")
    except Exception:
        valid_until = ""

    left_info = [
        f"Nội dung xuất: {issue.notes or ''}",
        f"Kho giao hàng: {issue.from_warehouse_name or ''}",
        f"Người nhận: {issue.recipient or ''}",
        f"Hàng do: {issue.transporter or ''}",
    ]
    right_info = [
        f"Đơn vị nhận: {issue.to_warehouse_name or ''}",
        f"Có giá trị đến ngày: {valid_until}",
        "",
        "",
    ]
    for ri, (ltext, rtext) in enumerate(zip(left_info, right_info), start=4):
        sh(ri, 18)
        ws.merge_cells(f"A{ri}:D{ri}")
        c = ws.cell(row=ri, column=1, value=ltext)
        c.font = Font(name=TN, size=11); c.alignment = aln("left")
        if rtext:
            ws.merge_cells(f"E{ri}:H{ri}")
            c = ws.cell(row=ri, column=5, value=rtext)
            c.font = Font(name=TN, size=11); c.alignment = aln("left")

    # ── Row 8: blank ──────────────────────────────────────────────────────────
    sh(8, 8)

    # ── Rows 9-10: table header ───────────────────────────────────────────────
    hf   = Font(name=TN, size=11, bold=True)
    ha   = Alignment(horizontal="center", vertical="center", wrap_text=True)
    hfill = PatternFill("solid", fgColor="F2F2F2")
    sh(9, 22); sh(10, 20)
    for ci, val in enumerate(["TT","Tên hàng","ĐVT","Số lượng","","Đơn giá","Thành tiền","Ghi chú"], 1):
        c = ws.cell(row=9, column=ci, value=val)
        c.font = hf; c.alignment = ha; c.border = bdr; c.fill = hfill
    for ci, val in enumerate(["","","","Kế hoạch","Thực hiện","","",""], 1):
        c = ws.cell(row=10, column=ci, value=val)
        c.font = hf; c.alignment = ha; c.border = bdr; c.fill = hfill
    ws.merge_cells("A9:A10"); ws.merge_cells("B9:B10"); ws.merge_cells("C9:C10")
    ws.merge_cells("D9:E9")
    ws.merge_cells("F9:F10"); ws.merge_cells("G9:G10"); ws.merge_cells("H9:H10")

    # ── Data rows ─────────────────────────────────────────────────────────────
    grand_total = 0.0
    DR = 11
    for i, line in enumerate(lines):
        r = DR + i
        sh(r, 20)
        is_dc   = getattr(line, "is_shared", False)
        name    = f"{line.item_name} (DC)" if is_dc else line.item_name
        if is_dc:
            price = 0.0; lt = 0.0
        else:
            price = getattr(line, "unit_price", 0.0) or 0.0
            lt    = line.quantity * price
            grand_total += lt

        vals = [i + 1, name, line.unit_of_measure, line.quantity, "",
                price if (price and not is_dc) else "",
                lt    if (lt    and not is_dc) else "",
                getattr(line, "notes", "") or ""]
        alns = [aln("center"), aln("left"), aln("center"), aln("center"), aln("center"),
                aln("right"), aln("right"), aln("left")]
        for ci, (val, al) in enumerate(zip(vals, alns), 1):
            c = ws.cell(row=r, column=ci, value=val)
            c.font = Font(name=TN, size=11); c.alignment = al; c.border = bdr
            if ci in (6, 7) and isinstance(val, (int, float)) and val:
                c.number_format = "#,##0"

    # ── Tổng cộng row ─────────────────────────────────────────────────────────
    tr = DR + len(lines)
    sh(tr, 22)
    ws.merge_cells(f"A{tr}:F{tr}")
    c = ws.cell(row=tr, column=1, value="Tổng cộng")
    c.font = Font(name=TN, size=11, bold=True); c.alignment = aln("center"); c.border = bdr
    c = ws.cell(row=tr, column=7, value=grand_total if grand_total else "")
    c.font = Font(name=TN, size=11, bold=True); c.alignment = aln("right"); c.border = bdr
    if grand_total:
        c.number_format = "#,##0"
    ws.cell(row=tr, column=8).border = bdr

    # ── Footer ────────────────────────────────────────────────────────────────
    f1 = tr + 1; sh(f1, 18)
    ws.merge_cells(f"A{f1}:H{f1}")
    c = ws.cell(row=f1, column=1, value=f"Tổng: {len(lines):02d} Khoản")
    c.font = Font(name=TN, size=11); c.alignment = aln("left")

    f2 = tr + 2; sh(f2, 18)
    ws.merge_cells(f"A{f2}:H{f2}")
    c = ws.cell(row=f2, column=1,
                value=f"Thành tiền: {so_thanh_chu(int(grand_total))}.")
    c.font = Font(name=TN, size=11); c.alignment = aln("left")

    dy, mo, yr = _parse_date(issue.transaction_date or "")
    dr = tr + 3; sh(dr, 18)
    ws.merge_cells(f"E{dr}:H{dr}")
    c = ws.cell(row=dr, column=5,
                value=f"Ngày {dy}  tháng {mo}  năm {yr}")
    c.font = Font(name=TN, size=11, italic=True); c.alignment = aln("center")

    sh(dr + 1, 14)

    sr = dr + 2; sh(sr, 18)
    sigs = [("A", "A", "NGƯỜI LẬP PHIẾU"), ("B", "C", "NGƯỜI GIAO"),
            ("D", "E", "NGƯỜI NHẬN"), ("F", "G", "PHÒNG QUÂN NHU"), ("H", "H", "THỦ TRƯỞNG")]
    for cs, ce, label in sigs:
        if cs != ce:
            ws.merge_cells(f"{cs}{sr}:{ce}{sr}")
        ci = ord(cs) - ord("A") + 1
        c = ws.cell(row=sr, column=ci, value=label)
        c.font = Font(name=TN, size=11, bold=True); c.alignment = aln("center")

    try:
        wb.save(path)
        QMessageBox.information(parent, "Xuất Excel",
                                f"Đã lưu phiếu xuất kho:\n{path}")
    except Exception as exc:
        QMessageBox.warning(parent, "Lỗi", f"Không thể lưu file:\n{exc}")


def export_dieu_chinh(parent, transfer, lines) -> None:
    if not _ensure_docx(parent):
        return
    from PyQt6.QtWidgets import QMessageBox
    from docx.shared import Pt

    path = _save_path(parent, transfer.reference_number, "PhieuDieuChinh")
    if not path:
        return

    doc = _new_doc()
    h   = _helpers()
    _r, _fp, _ct, _set_bdr, _borders, L, C, R = h

    _title_section(doc, h, "PHIẾU ĐIỀU CHỈNH", transfer.reference_number)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    left = [
        ("Xuất từ kho", transfer.from_warehouse_name),
        ("Nhập vào kho", transfer.to_warehouse_name),
    ]
    if transfer.transporter:
        left.append(("Vận chuyển", transfer.transporter))

    _info_section(doc, h, left, transfer.notes)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    _table_simple(doc, h, lines, show_quality=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    sig = ["NGƯỜI LẬP PHIẾU", "NGƯỜI GIAO", "NGƯỜI NHẬN", "PHÒNG KỸ THUẬT", "THỦ TRƯỞNG"]
    _footer_section(doc, h, len(lines), 0, transfer.transaction_date, False, sig)

    doc.save(path)
    QMessageBox.information(parent, "Xuất Word", f"Đã lưu:\n{path}")
